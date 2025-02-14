# app/admin.py
from fastapi import APIRouter, Depends, HTTPException, File, UploadFile, Form, Query
from fastapi.security import HTTPBasic, HTTPBasicCredentials
from fastapi.responses import HTMLResponse
import secrets
from pathlib import Path
import os
import json
from typing import List, Optional
from io import BytesIO
import PyPDF2
import docx
from datetime import datetime
import logging

# Import our new metadata generator
from .metadata_generator import EnhancedMetadataGenerator

logger = logging.getLogger(__name__)

# Initialize the metadata generator
metadata_generator = EnhancedMetadataGenerator()

class MetadataExtractor:
    @staticmethod
    def extract_pdf_metadata(file_bytes: bytes) -> dict:
        metadata = {}
        try:
            pdf_file = BytesIO(file_bytes)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            # Extract basic PDF info
            metadata['page_count'] = len(pdf_reader.pages)
            metadata['title'] = pdf_reader.metadata.get('/Title', '')
            metadata['author'] = pdf_reader.metadata.get('/Author', '')
            metadata['creation_date'] = pdf_reader.metadata.get('/CreationDate', '')
            
            # Extract text from first page for topic inference
            first_page_text = pdf_reader.pages[0].extract_text()
            metadata['first_page_preview'] = first_page_text[:200] if first_page_text else ''
            
            # Infer document type and category
            metadata['document_type'] = 'pdf'
            metadata['category'] = MetadataExtractor.infer_category(first_page_text)
            
        except Exception as e:
            logger.error(f"Error extracting PDF metadata: {str(e)}")
        return metadata

    @staticmethod
    def extract_docx_metadata(file_bytes: bytes) -> dict:
        metadata = {}
        try:
            doc = docx.Document(BytesIO(file_bytes))
            
            # Extract basic DOCX info
            metadata['page_count'] = len(doc.paragraphs)
            core_properties = doc.core_properties
            metadata['title'] = core_properties.title or ''
            metadata['author'] = core_properties.author or ''
            metadata['created'] = str(core_properties.created) if core_properties.created else ''
            
            # Extract text from first paragraph for topic inference
            first_text = doc.paragraphs[0].text if doc.paragraphs else ''
            metadata['first_page_preview'] = first_text[:200]
            
            # Infer document type and category
            metadata['document_type'] = 'docx'
            metadata['category'] = MetadataExtractor.infer_category(first_text)
            
        except Exception as e:
            logger.error(f"Error extracting DOCX metadata: {str(e)}")
        return metadata

    @staticmethod
    def extract_txt_metadata(content: str) -> dict:
        metadata = {}
        try:
            # Basic text analysis
            lines = content.split('\n')
            metadata['line_count'] = len(lines)
            metadata['first_page_preview'] = content[:200]
            
            # Infer document type and category
            metadata['document_type'] = 'txt'
            metadata['category'] = MetadataExtractor.infer_category(content)
            
        except Exception as e:
            logger.error(f"Error extracting TXT metadata: {str(e)}")
        return metadata

    @staticmethod
    def infer_category(text: str) -> str:
        """Infer document category based on content analysis"""
        text = text.lower()
        
        # Define category keywords
        categories = {
            'pricing': ['price', 'cost', 'fee', 'payment', 'rates'],
            'policies': ['policy', 'rule', 'guideline', 'requirement', 'terms'],
            'venue_info': ['capacity', 'space', 'facility', 'amenity', 'feature'],
            'catering': ['food', 'beverage', 'menu', 'catering', 'dining'],
            'contracts': ['contract', 'agreement', 'legal', 'terms', 'conditions'],
            'schedules': ['schedule', 'timeline', 'date', 'time', 'availability'],
            'vendor_info': ['vendor', 'supplier', 'service', 'provider', 'partner']
        }
        
        # Count keyword matches for each category
        category_scores = {cat: sum(1 for kw in kws if kw in text) 
                         for cat, kws in categories.items()}
        
        # Return category with highest score, or 'general' if no matches
        max_score = max(category_scores.values())
        if max_score > 0:
            return max(category_scores.items(), key=lambda x: x[1])[0]
        return 'general'

class DocumentProcessor:
    """Handles document processing and content extraction"""
    
    @staticmethod
    def extract_content(file_bytes: bytes, filename: str) -> tuple[str, dict]:
        """Extract content and basic file info from uploaded file"""
        file_info = {
            'filename': filename,
            'size': len(file_bytes),
            'upload_date': datetime.now().isoformat()
        }
        
        if filename.lower().endswith('.pdf'):
            content = DocumentProcessor._extract_pdf_content(file_bytes)
            file_info['content_type'] = 'application/pdf'
        elif filename.lower().endswith('.docx'):
            content = DocumentProcessor._extract_docx_content(file_bytes)
            file_info['content_type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        elif filename.lower().endswith('.txt'):
            content = file_bytes.decode('utf-8')
            file_info['content_type'] = 'text/plain'
        else:
            raise ValueError(f"Unsupported file type: {filename}")
            
        return content, file_info

    @staticmethod
    def _extract_pdf_content(file_bytes: bytes) -> str:
        """Extract text content from PDF"""
        try:
            pdf_file = BytesIO(file_bytes)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            # Extract text from all pages
            content = []
            for page in pdf_reader.pages:
                content.append(page.extract_text())
                
            return '\n'.join(content)
        except Exception as e:
            logger.error(f"Error extracting PDF content: {str(e)}")
            raise

    @staticmethod
    def _extract_docx_content(file_bytes: bytes) -> str:
        """Extract text content from DOCX"""
        try:
            doc = docx.Document(BytesIO(file_bytes))
            content = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                content.append(paragraph.text)
                
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        content.append(cell.text)
                        
            return '\n'.join(content)
        except Exception as e:
            logger.error(f"Error extracting DOCX content: {str(e)}")
            raise

# Router setup
router = APIRouter()
security = HTTPBasic()

# Auth settings
ADMIN_USERNAME = os.getenv("ADMIN_USERNAME", "admin")
ADMIN_PASSWORD = os.getenv("ADMIN_PASSWORD", "changeme")

def get_current_admin(credentials: HTTPBasicCredentials = Depends(security)):
    is_correct_username = secrets.compare_digest(credentials.username, ADMIN_USERNAME)
    is_correct_password = secrets.compare_digest(credentials.password, ADMIN_PASSWORD)
    
    if not (is_correct_username and is_correct_password):
        raise HTTPException(
            status_code="HTTP_401_UNAUTHORIZED",
            detail="Incorrect username or password",
            headers={"WWW-Authenticate": "Basic"},
        )
    return credentials.username

@router.get("/admin", response_class=HTMLResponse)
async def admin_interface(_: str = Depends(get_current_admin)):
    """Serve the admin interface HTML"""
    html_path = Path(__file__).parent / "templates" / "admin.html"
    with open(html_path, "r") as f:
        return f.read()

@router.get("/admin/document/{doc_id}/metadata")
async def get_document_metadata(
    doc_id: str,
    _: str = Depends(get_current_admin)
):
    """Get detailed metadata for a specific document"""
    try:
        from app.main import llm_manager
        
        # Get document
        doc = llm_manager.get_document(doc_id)
        if not doc:
            raise HTTPException(status_code=404, detail="Document not found")
            
        # Generate metadata summary
        metadata_summary = metadata_generator.export_metadata_summary(doc['metadata'])
        
        return {
            'metadata': doc['metadata'],
            'summary': metadata_summary
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error getting document metadata: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))
    
@router.put("/admin/document/{doc_id}/metadata")
async def update_document_metadata(
    doc_id: str,
    updates: dict,
    _: str = Depends(get_current_admin)
):
    """Update metadata for a specific document"""
    try:
        from app.main import llm_manager
        
        # Get existing document
        doc = llm_manager.get_document(doc_id)
        if not doc:
            raise HTTPException(status_code=404, detail="Document not found")
            
        # Update metadata
        updated_metadata = metadata_generator.update_metadata(
            content=doc['text'],
            existing_metadata=doc['metadata'],
            update_fields=updates.get('fields')
        )
        
        # Save updated metadata
        success = llm_manager.update_document_metadata(doc_id, updated_metadata)
        if not success:
            raise HTTPException(status_code=500, detail="Failed to update metadata")
            
        return {
            'message': 'Metadata updated successfully',
            'metadata': updated_metadata
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error updating document metadata: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/admin/documents")
async def list_documents(
    _: str = Depends(get_current_admin),
    page: int = Query(1, ge=1),
    search: str = Query(""),
    page_size: int = Query(10, ge=1, le=100),
    category: Optional[str] = Query(None),
    content_type: Optional[str] = Query(None),
    date_from: Optional[str] = Query(None),
    date_to: Optional[str] = Query(None)
):
    """List documents with enhanced filtering and search"""
    try:
        from app.main import llm_manager
        
        # Get base document list
        documents = llm_manager.list_documents(
            search_term=search,
            page=page,
            page_size=page_size
        )
        
        # Apply additional filters if specified
        filtered_docs = []
        for doc in documents['documents']:
            # Skip if document doesn't match category filter
            if category and category not in doc.get('metadata', {}).get('categories', []):
                continue
                
            # Skip if document doesn't match content type filter
            if content_type and content_type != doc.get('metadata', {}).get('content_type'):
                continue
                
            # Skip if document is outside date range
            doc_date = doc.get('metadata', {}).get('timestamp')
            if doc_date:
                if date_from and doc_date < date_from:
                    continue
                if date_to and doc_date > date_to:
                    continue
                    
            # Document passed all filters
            filtered_docs.append(doc)
        
        # Update response with filtered documents
        documents['documents'] = filtered_docs
        documents['total'] = len(filtered_docs)
        
        return documents
        
    except Exception as e:
        logger.error(f"Error listing documents: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/admin/upload")
async def upload_files(
    file: UploadFile = File(...),
    metadata_json: Optional[str] = Form(None),
    _: str = Depends(get_current_admin)
):
    """Handle file upload with enhanced metadata extraction"""
    try:
        from app.main import llm_manager
        
        # Read file content
        content = await file.read()
        
        # Extract content and basic file info
        text_content, file_info = DocumentProcessor.extract_content(content, file.filename)
        
        # Parse user-provided metadata if any
        user_metadata = {}
        if metadata_json:
            try:
                user_metadata = json.loads(metadata_json)
            except json.JSONDecodeError:
                logger.warning("Invalid metadata JSON provided")
        
        # Generate enhanced metadata
        metadata = metadata_generator.generate_metadata(
            content=text_content,
            file_info=file_info,
            existing_metadata=user_metadata
        )
        
        # Validate metadata
        validation_issues = metadata_generator.validate_metadata(metadata)
        if validation_issues:
            logger.warning(f"Metadata validation issues: {validation_issues}")
            metadata['validation_warnings'] = validation_issues
        
        # Add to knowledge base
        doc_id = llm_manager.add_to_knowledge_base(
            text=text_content,
            metadata=metadata
        )
        
        # Generate a human-readable summary for the response
        metadata_summary = metadata_generator.export_metadata_summary(metadata)

        return {
            'message': 'File uploaded successfully',
            'doc_id': doc_id,
            'metadata_summary': metadata_summary,
            'metadata': metadata
        }
        
    except ValueError as ve:
        logger.error(f"Validation error during file upload: {str(ve)}")
        raise HTTPException(status_code=400, detail=str(ve))
    except Exception as e:
        logger.error(f"Error during file upload: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@router.delete("/admin/documents/{doc_id}")
async def delete_document(
    doc_id: str,
    _: str = Depends(get_current_admin)
):
    """Delete a document"""
    try:
        from app.main import llm_manager
        success = llm_manager.delete_document(doc_id)
        if success:
            return {"message": "Document deleted successfully"}
        raise HTTPException(status_code=404, detail="Document not found")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))